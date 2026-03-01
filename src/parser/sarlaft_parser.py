"""
SARLAFT Document Parser & Chunker
==================================
Parses the CBJ SFC SARLAFT chapter (Parte I, Título IV, Capítulo IV)
from .docx format into hierarchically-structured chunks with metadata.

Designed for RAG ingestion. Preserves the legal numbering hierarchy
(up to 10 levels deep) and injects parent context into each chunk.

Usage:
    python sarlaft_parser.py <input.docx> <output_dir>

Output:
    - chunks.json: All chunks with metadata
    - chunks_preview.txt: Human-readable preview for validation
    - parsing_report.txt: Stats and diagnostics
"""

import json
import re
import sys
import os
from dataclasses import dataclass, field, asdict
from typing import Optional
from docx import Document


# ============================================================
# Data structures
# ============================================================

@dataclass
class Section:
    """Represents a numbered section in the document."""
    numeral: str            # e.g. "4.2.2.2.1.1."
    title: str              # Short title if bold heading, else first ~80 chars
    full_text: str          # Complete text content (without children)
    depth: int              # Number of dots = hierarchy depth
    paragraph_index: int    # Position in original document
    is_bold_heading: bool   # Whether this was a bold heading paragraph
    children: list = field(default_factory=list)


@dataclass
class Chunk:
    """A chunk ready for embedding and storage in vector DB."""
    chunk_id: str
    source: str
    section_id: str
    section_title: str
    hierarchy_path: str         # Full breadcrumb path
    hierarchy_path_ids: str     # Numeral-only path: "4. > 4.2. > 4.2.2."
    depth: int
    content: str                # The actual text with context prefix
    raw_content: str            # Text without context prefix
    char_count: int
    last_updated: str
    parent_section: str
    chunk_type: str             # "section", "definition", "intro", "concatenated_list"


# ============================================================
# Constants
# ============================================================

SOURCE_LABEL = "CBJ SFC, Parte I, Título IV, Capítulo IV — SARLAFT"
LAST_UPDATED = "2025-06"
MAX_CHUNK_CHARS = 3000  # ~750 tokens. If exceeded, subdivide.
MIN_CHUNK_CHARS = 100   # Chunks below this get merged with parent.

# Matches numbered sections: "1.", "1.1.", "4.2.2.2.1.1.1.3.2.1."
NUMERAL_RE = re.compile(r'^(\d+\.(?:\d+\.)*)\s*(.*)', re.DOTALL)


# ============================================================
# Parsing
# ============================================================

def parse_docx(filepath: str) -> tuple[list[dict], str]:
    """
    Parse the .docx into a list of raw paragraph records.
    Returns (paragraphs, footer_text).
    """
    doc = Document(filepath)

    # Extract footer for reference
    footer_text = ""
    for section in doc.sections:
        for p in section.footer.paragraphs:
            if p.text.strip():
                footer_text += p.text.strip() + " "

    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        is_bold = any(run.bold for run in p.runs if run.bold is True)

        paragraphs.append({
            "index": i,
            "text": text,
            "is_bold": is_bold,
        })

    return paragraphs, footer_text.strip()


def extract_intro(paragraphs: list[dict]) -> tuple[Optional[dict], int]:
    """
    Extract the introductory text before numeral 1 (Consideraciones generales).
    Returns (intro_record, start_index_for_numerals).
    """
    intro_parts = []
    start_idx = 0

    for i, p in enumerate(paragraphs):
        match = NUMERAL_RE.match(p["text"])
        if match:
            start_idx = i
            break
        # Skip the document header lines (first 4 bold lines)
        if p["is_bold"] and i < 5:
            continue
        intro_parts.append(p["text"])

    if intro_parts:
        intro_text = "\n\n".join(intro_parts)
        return {
            "numeral": "0.",
            "title": "Consideraciones generales",
            "text": intro_text,
            "depth": 1,
            "index": 0,
            "is_bold": False,
        }, start_idx

    return None, start_idx


def build_section_list(paragraphs: list[dict], start_idx: int) -> list[Section]:
    """
    Build a flat list of Section objects from the numbered paragraphs.
    Non-numbered paragraphs are attached to the preceding numbered section.
    """
    sections = []
    current_section = None

    for p in paragraphs[start_idx:]:
        text = p["text"]
        match = NUMERAL_RE.match(text)

        if match:
            # Save previous section
            if current_section:
                sections.append(current_section)

            numeral = match.group(1)
            rest = match.group(2).strip()
            depth = numeral.count(".")

            # Determine title
            if p["is_bold"]:
                title = rest[:120] if rest else numeral
            else:
                # For definitions and similar, extract the term before ":"
                colon_pos = rest.find(":")
                if colon_pos > 0 and colon_pos < 80:
                    title = rest[:colon_pos].strip()
                elif len(rest) < 80:
                    title = rest
                else:
                    title = rest[:80] + "..."

            current_section = Section(
                numeral=numeral,
                title=title,
                full_text=text,
                depth=depth,
                paragraph_index=p["index"],
                is_bold_heading=p["is_bold"],
            )
        else:
            # Non-numbered paragraph: append to current section
            if current_section:
                current_section.full_text += "\n\n" + text
            # If no current section yet, this is orphan text (shouldn't happen after intro)

    # Don't forget the last section
    if current_section:
        sections.append(current_section)

    return sections


def build_hierarchy_map(sections: list[Section]) -> dict[str, Section]:
    """Build a lookup map from numeral -> Section."""
    return {s.numeral: s for s in sections}


def get_parent_numeral(numeral: str) -> Optional[str]:
    """
    Get parent numeral. E.g.:
        "4.2.2." -> "4.2."
        "4." -> None
        "1.15.1." -> "1.15."
    """
    parts = numeral.rstrip(".").split(".")
    if len(parts) <= 1:
        return None
    return ".".join(parts[:-1]) + "."


def build_hierarchy_path(numeral: str, section_map: dict[str, Section]) -> tuple[str, str]:
    """
    Build the full breadcrumb path for a numeral.
    Returns (human_readable_path, numeral_only_path).
    """
    parts = []
    numeral_parts = []
    current = numeral

    while current:
        if current in section_map:
            s = section_map[current]
            # Use short title for path
            short_title = s.title[:60]
            parts.append(f"{s.numeral} {short_title}")
            numeral_parts.append(s.numeral)
        current = get_parent_numeral(current)

    parts.reverse()
    numeral_parts.reverse()

    return " > ".join(parts), " > ".join(numeral_parts)


# ============================================================
# Chunking
# ============================================================

def should_chunk_at_level(section: Section, sections: list[Section], section_map: dict[str, Section]) -> bool:
    """
    Determine if a section should be a chunk boundary.
    
    Strategy: chunk at the level where we get semantically complete units
    with enough context. Merge small leaf nodes into their parents.
    
    A section is a chunk boundary if:
    1. It's depth 1-3 (always: these are major structural units)
    2. It's depth 4-5 AND has substantial content (>200 chars own text)
       OR has children that accumulate to >300 chars
    3. It's depth 6+ only if it has >500 chars of own text
    """
    # Major sections: always chunk
    if section.depth <= 3:
        return True
    
    # Mid-level: chunk if substantial
    if section.depth <= 5:
        # Check own text length
        if len(section.full_text) > 200:
            return True
        # Check if it has children (meaning it's a grouping node)
        children = [s for s in sections if s.numeral.startswith(section.numeral) 
                     and s.numeral != section.numeral]
        if children:
            total = sum(len(c.full_text) for c in children) + len(section.full_text)
            if total > 300:
                return True
        return False
    
    # Deep levels: only chunk if they have a lot of text
    if len(section.full_text) > 500:
        return True
    
    return False


def collect_children_text(
    numeral: str,
    sections: list[Section],
    section_map: dict[str, Section],
    max_depth: Optional[int] = None,
) -> list[Section]:
    """
    Collect all direct and indirect children of a numeral.
    """
    children = []
    for s in sections:
        if s.numeral == numeral:
            continue
        if s.numeral.startswith(numeral) and s.numeral != numeral:
            if max_depth is None or s.depth <= max_depth:
                children.append(s)
    return children


def create_chunks(sections: list[Section], intro_record: Optional[dict]) -> list[Chunk]:
    """
    Create chunks from sections using hierarchical grouping strategy.
    """
    section_map = build_hierarchy_map(sections)
    chunks = []
    consumed = set()  # Track which sections have been included in a chunk

    # Intro chunk
    if intro_record:
        chunks.append(Chunk(
            chunk_id="SARLAFT_intro",
            source=SOURCE_LABEL,
            section_id="0.",
            section_title="Consideraciones generales",
            hierarchy_path="Consideraciones generales",
            hierarchy_path_ids="0.",
            depth=0,
            content=f"[Fuente: {SOURCE_LABEL}]\n[Sección: Consideraciones generales]\n\n{intro_record['text']}",
            raw_content=intro_record["text"],
            char_count=len(intro_record["text"]),
            last_updated=LAST_UPDATED,
            parent_section="",
            chunk_type="intro",
        ))

    # Process sections in order
    for section in sections:
        if section.numeral in consumed:
            continue

        if not should_chunk_at_level(section, sections, section_map):
            # Will be consumed by parent
            continue

        # Collect children that won't be their own chunks
        children_to_include = []
        children = collect_children_text(section.numeral, sections, section_map)

        for child in children:
            if child.numeral in consumed:
                continue
            if should_chunk_at_level(child, sections, section_map):
                # This child will be its own chunk, don't include
                continue
            children_to_include.append(child)

        # Build chunk text
        text_parts = [section.full_text]
        consumed.add(section.numeral)

        for child in children_to_include:
            text_parts.append(child.full_text)
            consumed.add(child.numeral)

        raw_content = "\n\n".join(text_parts)

        # Check if chunk is too large and needs splitting
        if len(raw_content) > MAX_CHUNK_CHARS and children_to_include:
            # Split: make the parent its own chunk, children become separate chunks
            _create_single_chunk(section, section.full_text, section_map, chunks)
            consumed.discard(section.numeral)  # re-add as its own
            consumed.add(section.numeral)

            for child in children_to_include:
                _create_single_chunk(child, child.full_text, section_map, chunks)
        else:
            # Single chunk with children rolled in
            _create_single_chunk(section, raw_content, section_map, chunks)

    # Catch any unconsumed sections (shouldn't happen but safety net)
    for section in sections:
        if section.numeral not in consumed:
            _create_single_chunk(section, section.full_text, section_map, chunks)
            consumed.add(section.numeral)

    return chunks


def _create_single_chunk(
    section: Section,
    raw_content: str,
    section_map: dict[str, Section],
    chunks: list[Chunk],
):
    """Create a single Chunk object and append to chunks list."""
    hierarchy_path, hierarchy_path_ids = build_hierarchy_path(section.numeral, section_map)
    parent = get_parent_numeral(section.numeral) or ""

    # Build context prefix
    context_prefix = f"[Fuente: {SOURCE_LABEL}]\n[Sección: {hierarchy_path}]\n\n"
    content = context_prefix + raw_content

    # Determine chunk type
    if section.numeral.startswith("1.") and section.depth == 2:
        chunk_type = "definition"
    elif section.depth <= 2:
        chunk_type = "section"
    else:
        chunk_type = "section"

    chunk_id = f"SARLAFT_{section.numeral.rstrip('.')}"

    chunks.append(Chunk(
        chunk_id=chunk_id,
        source=SOURCE_LABEL,
        section_id=section.numeral,
        section_title=section.title,
        hierarchy_path=hierarchy_path,
        hierarchy_path_ids=hierarchy_path_ids,
        depth=section.depth,
        content=content,
        raw_content=raw_content,
        char_count=len(raw_content),
        last_updated=LAST_UPDATED,
        parent_section=parent,
        chunk_type=chunk_type,
    ))


def merge_small_chunks(chunks: list[Chunk], min_chars: int = 100) -> list[Chunk]:
    """
    Merge chunks smaller than min_chars into their previous sibling chunk
    (same depth, same parent). Falls back to parent chunk if no sibling exists.
    This avoids merging leaf definitions into structural parent nodes.
    """
    chunk_map = {c.section_id: c for c in chunks}
    # Build ordered list of section_ids for sibling lookup
    ordered_ids = [c.section_id for c in chunks]
    to_remove = set()
    
    # Process from deepest to shallowest
    sorted_chunks = sorted(chunks, key=lambda c: -c.depth)
    
    # Build set of section_ids that are parents (have children in the chunk list)
    all_parents = set()
    for chunk in chunks:
        if chunk.parent_section:
            all_parents.add(chunk.parent_section)
    
    for chunk in sorted_chunks:
        if chunk.char_count >= min_chars:
            continue
        if chunk.chunk_type == "intro":
            continue
        if chunk.chunk_type == "definition":
            continue
        # Never merge a section that has children — it's a structural heading
        if chunk.section_id in all_parents:
            continue
        
        # Try to find previous sibling (same parent, same depth)
        merged = False
        chunk_idx = ordered_ids.index(chunk.section_id) if chunk.section_id in ordered_ids else -1
        
        if chunk_idx > 0:
            # Walk backwards to find a sibling
            for j in range(chunk_idx - 1, -1, -1):
                candidate = chunk_map.get(ordered_ids[j])
                if candidate is None or candidate.section_id in to_remove:
                    continue
                # Same parent = same depth and shares parent prefix
                if (candidate.depth == chunk.depth 
                    and candidate.parent_section == chunk.parent_section):
                    # Merge into previous sibling
                    candidate.raw_content += "\n\n" + chunk.raw_content
                    candidate.content = (
                        f"[Fuente: {SOURCE_LABEL}]\n"
                        f"[Sección: {candidate.hierarchy_path}]\n\n"
                        f"{candidate.raw_content}"
                    )
                    candidate.char_count = len(candidate.raw_content)
                    to_remove.add(chunk.section_id)
                    merged = True
                    break
        
        # Fallback: merge into parent if no sibling found
        if not merged and chunk.parent_section and chunk.parent_section in chunk_map:
            parent = chunk_map[chunk.parent_section]
            if parent.section_id not in to_remove:
                parent.raw_content += "\n\n" + chunk.raw_content
                parent.content = (
                    f"[Fuente: {SOURCE_LABEL}]\n"
                    f"[Sección: {parent.hierarchy_path}]\n\n"
                    f"{parent.raw_content}"
                )
                parent.char_count = len(parent.raw_content)
                to_remove.add(chunk.section_id)
    
    return [c for c in chunks if c.section_id not in to_remove]


def split_oversized_chunks(
    chunks: list[Chunk],
    section_map: dict[str, Section],
    max_chars: int = 3000,
) -> list[Chunk]:
    """
    Split chunks larger than max_chars into sub-chunks.
    Splits on double newlines (paragraph boundaries) while preserving
    the context prefix.
    """
    result = []
    
    for chunk in chunks:
        if chunk.char_count <= max_chars:
            result.append(chunk)
            continue
        
        # Split raw content on paragraph boundaries
        paragraphs = chunk.raw_content.split("\n\n")
        
        sub_chunks = []
        current_text = ""
        part_num = 0
        
        for para in paragraphs:
            if current_text and len(current_text) + len(para) + 2 > max_chars:
                # Emit current sub-chunk
                part_num += 1
                sub_chunk = Chunk(
                    chunk_id=f"{chunk.chunk_id}_p{part_num}",
                    source=chunk.source,
                    section_id=f"{chunk.section_id}[p{part_num}]",
                    section_title=f"{chunk.section_title} (parte {part_num})",
                    hierarchy_path=chunk.hierarchy_path,
                    hierarchy_path_ids=chunk.hierarchy_path_ids,
                    depth=chunk.depth,
                    content=f"[Fuente: {SOURCE_LABEL}]\n[Sección: {chunk.hierarchy_path}]\n\n{current_text}",
                    raw_content=current_text,
                    char_count=len(current_text),
                    last_updated=chunk.last_updated,
                    parent_section=chunk.parent_section,
                    chunk_type=chunk.chunk_type,
                )
                sub_chunks.append(sub_chunk)
                current_text = para
            else:
                if current_text:
                    current_text += "\n\n" + para
                else:
                    current_text = para
        
        # Last sub-chunk
        if current_text:
            part_num += 1
            if part_num == 1:
                # Didn't actually split (single paragraph too long) — keep original
                result.append(chunk)
                continue
                
            sub_chunk = Chunk(
                chunk_id=f"{chunk.chunk_id}_p{part_num}",
                source=chunk.source,
                section_id=f"{chunk.section_id}[p{part_num}]",
                section_title=f"{chunk.section_title} (parte {part_num})",
                hierarchy_path=chunk.hierarchy_path,
                hierarchy_path_ids=chunk.hierarchy_path_ids,
                depth=chunk.depth,
                content=f"[Fuente: {SOURCE_LABEL}]\n[Sección: {chunk.hierarchy_path}]\n\n{current_text}",
                raw_content=current_text,
                char_count=len(current_text),
                last_updated=chunk.last_updated,
                parent_section=chunk.parent_section,
                chunk_type=chunk.chunk_type,
            )
            sub_chunks.append(sub_chunk)
        
        result.extend(sub_chunks)
    
    return result


# ============================================================
# Output
# ============================================================

def write_outputs(chunks: list[Chunk], output_dir: str):
    """Write chunks to JSON, preview, and report files."""
    os.makedirs(output_dir, exist_ok=True)

    # Sort chunks by document order (numeric sort on section_id parts)
    def sort_key(chunk):
        if chunk.section_id == "0.":
            return (0,)
        # Handle split chunks like "4.2.2.2.1.[p2]"
        clean_id = chunk.section_id.replace("[", ".").replace("]", "")
        parts = clean_id.rstrip(".").split(".")
        result = []
        for p in parts:
            if p.startswith("p"):
                result.append(int(p[1:]) + 10000)  # sort splits after parent
            else:
                try:
                    result.append(int(p))
                except ValueError:
                    result.append(0)
        return tuple(result)
    
    chunks = sorted(chunks, key=sort_key)

    # JSON output
    json_path = os.path.join(output_dir, "chunks.json")
    chunks_data = [asdict(c) for c in chunks]
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(chunks_data, f, ensure_ascii=False, indent=2)

    # Preview
    preview_path = os.path.join(output_dir, "chunks_preview.txt")
    with open(preview_path, "w", encoding="utf-8") as f:
        for c in chunks:
            f.write(f"{'='*80}\n")
            f.write(f"CHUNK: {c.chunk_id}\n")
            f.write(f"Section: {c.section_id} | Title: {c.section_title}\n")
            f.write(f"Path: {c.hierarchy_path}\n")
            f.write(f"Depth: {c.depth} | Chars: {c.char_count} | Type: {c.chunk_type}\n")
            f.write(f"Parent: {c.parent_section}\n")
            f.write(f"{'-'*80}\n")
            # Show first 500 chars of raw content
            preview = c.raw_content[:500]
            if len(c.raw_content) > 500:
                preview += "\n[...truncated...]"
            f.write(f"{preview}\n\n")

    # Report
    report_path = os.path.join(output_dir, "parsing_report.txt")
    char_counts = [c.char_count for c in chunks]
    depth_dist = {}
    type_dist = {}
    for c in chunks:
        depth_dist[c.depth] = depth_dist.get(c.depth, 0) + 1
        type_dist[c.chunk_type] = type_dist.get(c.chunk_type, 0) + 1

    with open(report_path, "w", encoding="utf-8") as f:
        f.write("SARLAFT PARSING REPORT\n")
        f.write(f"{'='*50}\n\n")
        f.write(f"Total chunks: {len(chunks)}\n")
        f.write(f"Total characters: {sum(char_counts):,}\n")
        f.write(f"Avg chunk size: {sum(char_counts)/len(char_counts):.0f} chars\n")
        f.write(f"Min chunk size: {min(char_counts)} chars\n")
        f.write(f"Max chunk size: {max(char_counts)} chars\n")
        f.write(f"Median chunk size: {sorted(char_counts)[len(char_counts)//2]} chars\n\n")

        f.write("Chunks by depth:\n")
        for d in sorted(depth_dist):
            f.write(f"  Depth {d}: {depth_dist[d]} chunks\n")

        f.write(f"\nChunks by type:\n")
        for t, count in sorted(type_dist.items()):
            f.write(f"  {t}: {count}\n")

        f.write(f"\nSize distribution:\n")
        brackets = [(0, 200), (200, 500), (500, 1000), (1000, 2000), (2000, 3000), (3000, 5000), (5000, 10000)]
        for lo, hi in brackets:
            count = sum(1 for c in char_counts if lo <= c < hi)
            if count > 0:
                f.write(f"  {lo}-{hi} chars: {count} chunks\n")

        # Flag potential issues
        f.write(f"\n--- Potential Issues ---\n")
        tiny = [c for c in chunks if c.char_count < MIN_CHUNK_CHARS]
        if tiny:
            f.write(f"Very small chunks (<{MIN_CHUNK_CHARS} chars): {len(tiny)}\n")
            for c in tiny[:10]:
                f.write(f"  {c.chunk_id}: {c.char_count} chars - {c.section_title[:60]}\n")

        huge = [c for c in chunks if c.char_count > MAX_CHUNK_CHARS]
        if huge:
            f.write(f"\nOversized chunks (>{MAX_CHUNK_CHARS} chars): {len(huge)}\n")
            for c in huge:
                f.write(f"  {c.chunk_id}: {c.char_count} chars - {c.section_title[:60]}\n")

    return json_path, preview_path, report_path


# ============================================================
# Main
# ============================================================

def main():
    if len(sys.argv) < 3:
        print(f"Usage: python {sys.argv[0]} <input.docx> <output_dir>")
        sys.exit(1)

    input_path = sys.argv[1]
    output_dir = sys.argv[2]

    print(f"Parsing: {input_path}")

    # Step 1: Parse docx
    paragraphs, footer = parse_docx(input_path)
    print(f"  Extracted {len(paragraphs)} non-empty paragraphs")
    print(f"  Footer: {footer[:100]}")

    # Step 2: Extract intro
    intro_record, start_idx = extract_intro(paragraphs)
    if intro_record:
        print(f"  Intro section found: {len(intro_record['text'])} chars")

    # Step 3: Build section list
    sections = build_section_list(paragraphs, start_idx)
    print(f"  Built {len(sections)} sections")

    # Step 4: Create chunks
    section_map = build_hierarchy_map(sections)
    chunks = create_chunks(sections, intro_record)
    print(f"  Created {len(chunks)} chunks")

    # Step 5: Post-process — merge very small chunks into their parents
    chunks = merge_small_chunks(chunks, min_chars=MIN_CHUNK_CHARS)
    print(f"  After merging small chunks: {len(chunks)} chunks")

    # Step 6: Post-process — split oversized chunks
    chunks = split_oversized_chunks(chunks, section_map, max_chars=MAX_CHUNK_CHARS)
    print(f"  After splitting oversized: {len(chunks)} chunks")

    # Step 7: Write outputs
    json_path, preview_path, report_path = write_outputs(chunks, output_dir)
    print(f"\nOutputs written to: {output_dir}/")
    print(f"  {json_path}")
    print(f"  {preview_path}")
    print(f"  {report_path}")


if __name__ == "__main__":
    main()
